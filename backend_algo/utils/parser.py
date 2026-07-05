import docx
import pdfplumber
import re


def parse_docx(file_path: str):
    doc = docx.Document(file_path)

    # 1. 提取所有纯文本段落
    paragraphs_text = []
    for p in doc.paragraphs:
        if p.text.strip():  # 过滤空行
            paragraphs_text.append(p.text.strip())

    # 2. 提取所有表格
    tables_data = []
    for table in doc.tables:
        table_matrix = []
        for row in table.rows:
            # 提取一行中每个单元格的文本
            row_text = [cell.text.strip() for cell in row.cells]
            table_matrix.append(row_text)
        tables_data.append(table_matrix)

    return paragraphs_text, tables_data



def parse_pdf(file_path: str):
    all_text = []
    all_tables = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # 1. 提取当前页的纯文本
            text = page.extract_text()
            if text:
                all_text.append(f"--- Page {page_num + 1} ---\n" + text)

            # 2. 提取当前页的表格（返回嵌套列表：[[行1], [行2]]）
            tables = page.extract_tables()
            for table in tables:
                if table:
                    all_tables.append(table)

    return all_text, all_tables


def table_to_markdown(table_matrix: list[list[str]]) -> str:
    """将二维列表表格转换为标准的 Markdown 格式字符串"""
    if not table_matrix or not table_matrix[0]:
        return ""

    md_lines = []
    # 1. 处理表头
    header = table_matrix[0]
    # 清理单元格内可能存在的回车，防止破坏 Markdown 表格结构
    clean_header = [str(cell).replace('\n', ' ') if cell else "" for cell in header]
    md_lines.append("| " + " | ".join(clean_header) + " |")

    # 2. 处理分割线
    md_lines.append("|" + "|".join(["---"] * len(header)) + "|")

    # 3. 处理数据行
    for row in table_matrix[1:]:
        clean_row = [str(cell).replace('\n', ' ') if cell else "" for cell in row]
        # 如果当前行比表头短，用空字符串补齐（防止表格解析错位）
        clean_row += [""] * (len(header) - len(clean_row))
        # 如果比表头长，截断（少见异常情况）
        clean_row = clean_row[:len(header)]
        md_lines.append("| " + " | ".join(clean_row) + " |")

    return "\n".join(md_lines)


def _is_noise_line(line: str) -> bool:
    """判断一行文本是否为噪声行（非标题内容）"""
    line = line.strip()
    if len(line) < 5:
        return True
    # 匹配纯数字/页码/卷期号等
    if re.match(r'^[\d\s\.\-/]+$', line):
        return True
    # 常见噪声模式: "Vol.XX", "第X卷", "2019 IEEE", "DOI:", "http"
    noise_patterns = [
        r'^(Vol|Volume|卷|期|No\.|Number)\s*\d',
        r'^\d{4}\s*(IEEE|ACM|Springer|CVPR|ICCV|ECCV|ACL|SIGMOD|VLDB)',
        r'^(DOI|doi|http|www\.)',
        r'^(摘\s*要|Abstract|Keywords|关键词)',
        r'^(第.+卷|第.+期)',
        r'^CHINA\s+PACKAGING',
    ]
    for pat in noise_patterns:
        if re.search(pat, line, re.IGNORECASE):
            return True
    return False


def _extract_title_from_text(text: str) -> str:
    """从首页文本中提取标题（取前几行中最长的有意义行）"""
    if not text:
        return ""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    # 只考虑前 8 行
    candidates = lines[:8]
    # 过滤掉噪声行
    candidates = [l for l in candidates if not _is_noise_line(l)]
    if not candidates:
        return ""
    # 合并连续行（论文标题可能跨行），取前2-3行拼接
    # 策略：取最长的单行，或前2行拼接（如果第2行不是噪声）
    best = candidates[0]
    if len(candidates) >= 2 and len(candidates[1]) > len(best) * 0.4:
        combined = candidates[0] + " " + candidates[1]
        if len(combined) <= 200:
            best = combined
    # 截断过长标题
    if len(best) > 200:
        best = best[:197] + "..."
    return best.strip()


def extract_title_from_pdf(file_path: str) -> str:
    """从 PDF 文件中提取标题：优先元数据，兜底首页文本"""
    try:
        with pdfplumber.open(file_path) as pdf:
            # 1. 优先使用元数据
            meta = pdf.metadata or {}
            title = (meta.get("Title") or "").strip()
            if title and len(title) > 3:
                return title

            # 2. 兜底：从首页文本提取
            if pdf.pages:
                first_text = pdf.pages[0].extract_text() or ""
                return _extract_title_from_text(first_text)
    except Exception:
        pass
    return ""


def extract_title_from_docx(file_path: str) -> str:
    """从 DOCX 文件中提取标题：优先 core_properties，兜底首段落"""
    try:
        doc = docx.Document(file_path)
        # 1. 优先元数据
        title = (doc.core_properties.title or "").strip()
        if title and len(title) > 3:
            return title

        # 2. 兜底：取前几个非空段落拼接
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return _extract_title_from_text("\n".join(paragraphs[:5]))
    except Exception:
        pass
    return ""


def extract_title(file_path: str, filename: str) -> str:
    """统一入口：根据文件类型提取标题，失败时返回文件名（去掉扩展名）"""
    if filename.lower().endswith(".pdf"):
        title = extract_title_from_pdf(file_path)
    elif filename.lower().endswith(".docx"):
        title = extract_title_from_docx(file_path)
    else:
        title = ""

    # 如果提取失败，用文件名（去扩展名）作为标题
    if not title:
        title = filename.rsplit(".", 1)[0] if "." in filename else filename
    return title